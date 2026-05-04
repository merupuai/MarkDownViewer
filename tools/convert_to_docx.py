#!/usr/bin/env python3
"""
Convert Markdown report to DOCX using python-docx
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def parse_markdown_to_docx(md_file_path, output_docx):
    """Parse markdown file and convert to DOCX"""

    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split('\n')
    i = 0
    in_table = False
    table_data = []
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines (but add spacing in document)
        if not line:
            if not in_table and not in_code_block:
                pass  # We'll handle spacing via paragraph styles
            i += 1
            continue

        # Code blocks
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # End of code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                in_code_block = False
            i += 1
            continue

        if in_code_block:
            code_lines.append(lines[i])  # Keep original formatting
            i += 1
            continue

        # Title (first H1)
        if line.startswith('# ') and i < 5:
            title_text = line[2:].strip()
            p = doc.add_heading(title_text, 0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(26, 26, 26)
            i += 1
            continue

        # H1
        if line.startswith('# '):
            h1_text = line[2:].strip()
            doc.add_page_break()
            p = doc.add_heading(h1_text, 1)
            for run in p.runs:
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(44, 62, 80)
            i += 1
            continue

        # H2
        if line.startswith('## '):
            h2_text = line[3:].strip()
            p = doc.add_heading(h2_text, 2)
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(52, 73, 94)
            i += 1
            continue

        # H3
        if line.startswith('### '):
            h3_text = line[4:].strip()
            p = doc.add_heading(h3_text, 3)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(52, 73, 94)
            i += 1
            continue

        # H4
        if line.startswith('#### '):
            h4_text = line[5:].strip()
            p = doc.add_paragraph(h4_text)
            p.style = 'Heading 4'
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(11)
            i += 1
            continue

        # Horizontal rule
        if line.startswith('---'):
            doc.add_paragraph('_' * 80)
            i += 1
            continue

        # Images
        if line.startswith('!['):
            # Extract image path and caption
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                caption = match.group(1)
                img_path = match.group(2)

                # Convert relative path to absolute
                if img_path.startswith('../charts/'):
                    img_path = '/workspace/charts/' + img_path[10:]
                elif img_path.startswith('charts/'):
                    img_path = '/workspace/' + img_path
                elif not img_path.startswith('/'):
                    img_path = '/workspace/' + img_path

                try:
                    doc.add_picture(img_path, width=Inches(6))
                    if caption:
                        p = doc.add_paragraph(caption)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(102, 102, 102)
                            run.italic = True
                except Exception as e:
                    print(f"Warning: Could not insert image {img_path}: {e}")

            i += 1
            continue

        # Tables
        if '|' in line and not in_table:
            # Start collecting table data
            in_table = True
            table_data = []

        if in_table:
            if '|' in line:
                # Skip separator line
                if '---' not in line:
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    table_data.append(cells)
                i += 1
                continue
            else:
                # End of table - create it
                if table_data:
                    num_cols = len(table_data[0])
                    table = doc.add_table(rows=len(table_data), cols=num_cols)
                    table.style = 'Light Grid Accent 1'

                    # Populate table
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            cell = table.rows[row_idx].cells[col_idx]
                            cell.text = cell_text

                            # Header row formatting
                            if row_idx == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                                        run.font.color.rgb = RGBColor(255, 255, 255)
                                cell._element.get_or_add_tcPr().append(
                                    parse_xml(r'<w:shd {} w:fill="3498db"/>'.format(nsdecls('w')))
                                )

                    doc.add_paragraph()  # Add spacing after table

                in_table = False
                table_data = []
                # Don't increment i, process this line again
                continue

        # Lists
        if line.startswith('- ') or line.startswith('* ') or re.match(r'^\d+\.', line):
            if line.startswith('- ') or line.startswith('* '):
                list_text = line[2:]
                # Process markdown formatting in list items
                list_text = process_markdown_formatting(list_text)
                p = doc.add_paragraph(list_text, style='List Bullet')
            elif re.match(r'^\d+\.', line):
                list_text = re.sub(r'^\d+\.\s*', '', line)
                list_text = process_markdown_formatting(list_text)
                p = doc.add_paragraph(list_text, style='List Number')

            for run in p.runs:
                run.font.size = Pt(10)

            i += 1
            continue

        # Regular paragraphs
        if line and not line.startswith('#'):
            # Process markdown formatting
            processed_line = process_markdown_formatting(line)

            p = doc.add_paragraph()
            add_formatted_text(p, line)

            for run in p.runs:
                run.font.size = Pt(10)

        i += 1

    # Save document
    doc.save(output_docx)
    print(f"DOCX generated successfully: {output_docx}")


def process_markdown_formatting(text):
    """Process basic markdown formatting markers"""
    # Bold
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)  # Bold italic
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
    # Links
    text = re.sub(r'\[(.*?)\]\((.*?)\)', r'\1', text)
    return text


def add_formatted_text(paragraph, text):
    """Add text to paragraph with markdown formatting"""
    # Handle bold/italic/links with runs
    parts = []
    current = ''
    i = 0

    while i < len(text):
        # Bold
        if text[i:i+2] == '**':
            if current:
                parts.append(('normal', current))
                current = ''
            # Find closing
            j = text.find('**', i+2)
            if j != -1:
                parts.append(('bold', text[i+2:j]))
                i = j + 2
                continue

        # Italic
        if text[i] == '*' and (i == 0 or text[i-1] != '*'):
            if current:
                parts.append(('normal', current))
                current = ''
            # Find closing
            j = text.find('*', i+1)
            if j != -1 and (j+1 >= len(text) or text[j+1] != '*'):
                parts.append(('italic', text[i+1:j]))
                i = j + 1
                continue

        # Links [text](url)
        if text[i] == '[':
            if current:
                parts.append(('normal', current))
                current = ''
            match = re.match(r'\[(.*?)\]\((.*?)\)', text[i:])
            if match:
                parts.append(('normal', match.group(1)))  # Just the link text
                i += len(match.group(0))
                continue

        current += text[i]
        i += 1

    if current:
        parts.append(('normal', current))

    # Add runs
    for style, txt in parts:
        run = paragraph.add_run(txt)
        if style == 'bold':
            run.bold = True
        elif style == 'italic':
            run.italic = True


def parse_xml(xml_string):
    """Helper to parse XML for table cell formatting"""
    from lxml import etree
    return etree.fromstring(xml_string)


def nsdecls(*prefixes):
    """Helper for XML namespace declarations"""
    return ' '.join(['xmlns:%s="%s"' % (p, qn(p+':foo').split('}')[0][1:])
                     for p in prefixes])


if __name__ == "__main__":
    md_file = "/workspace/docs/llm_gateway_research_report.md"
    output_docx = "/workspace/docs/llm_gateway_research_report.docx"

    parse_markdown_to_docx(md_file, output_docx)
